# auto-py-to-exe para compilar este script
# Asegúrate de tener instaladas las librerías necesarias
# pyinstaller --onefile --add-data "archivo.txt:." tu_script.py para linux

import os
import json
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from num2words import num2words
import re
from datetime import datetime
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import sys
from io import BytesIO

version = "v1.0.13"


# from docx2pdf import convert


def normaliza_fechas_realizacion(fecha_str):
    """
    Normaliza diferentes formatos de fechas a 'DD/MM/AA al DD/MM/AA'.
    Si no cumple el formato esperado, muestra un error y devuelve una cadena vacía.
    """
    if not fecha_str or fecha_str.lower() == 'nan':
        print(
            "Error en formato de fechas",
            "No se han encontrado fechas. Deben seguir el formato DD/MM/AA al DD/MM/AA"
        )
        return -1
    fecha_str = fecha_str.strip()
    # Buscar patrones tipo 'del XX/XX/XX al XX/XX/XX'
    match = re.search(r'(\d{1,2}/\d{1,2}/\d{2,4})\s*(?:-|al|a|hasta)\s*(\d{1,2}/\d{1,2}/\d{2,4})', fecha_str, re.IGNORECASE)
    if match:
        f1, f2 = match.group(1), match.group(2)
    else:
        # Buscar patrón 'XX/XX/XX-XX/XX/XX'
        match = re.search(r'(\d{1,2}/\d{1,2}/\d{2,4})\s*-\s*(\d{1,2}/\d{1,2}/\d{2,4})', fecha_str)
        if match:
            f1, f2 = match.group(1), match.group(2)
        else:
            # Buscar dos fechas separadas por espacio
            fechas = re.findall(r'\d{1,2}/\d{1,2}/\d{2,4}', fecha_str)
            if len(fechas) >= 2:
                f1, f2 = fechas[0], fechas[1]
            else:
                # Si solo hay una fecha, repetirla
                if fechas:
                    f1 = f2 = fechas[0]
                else:
                    # No se reconoce el formato, mostrar error y devolver cadena vacía
                    print(
                        "Error en format de dates",
                        "Les dates no estan bé, han de seguir el format DD/MM/AA al DD/MM/AA"
                    )
                    return -1
    # Formatear fechas a DD/MM/AA
    def corta_fecha(f):
        try:
            dt = datetime.strptime(f, "%d/%m/%Y")
            return dt.strftime("%d/%m/%y")
        except Exception:
            try:
                dt = datetime.strptime(f, "%d/%m/%y")
                return dt.strftime("%d/%m/%y")
            except Exception:
                return f
    
    # Comprobar la duración de las fechas

    # Normalizar formato antes de parsear
    try:
        fecha1 = datetime.strptime(f1, "%d/%m/%Y")
    except ValueError:
        fecha1 = datetime.strptime(f1, "%d/%m/%y")
    
    try:
        fecha2 = datetime.strptime(f2, "%d/%m/%Y")
    except ValueError:
        fecha2 = datetime.strptime(f2, "%d/%m/%y")
    duration = (fecha2 - fecha1).days

    
    if duration < 5:
        print(
            "Avís",
            f"La duració és de {duration} dies, que és menys de 5 dies. Desitja continuar?"
        )
    
    return f"{corta_fecha(f1)} al {corta_fecha(f2)}"

def extraer_datos_identificativos(nombre_archivo):
    df = pd.read_excel(nombre_archivo, header=None)
    etiquetas = [
        'CÓDIGO EDICIÓN / CODI EDICIÓ',
        'TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA',
        'FECHAS REALIZACIÓN / DATES REALITZACIÓ',
        'MODALIDAD/MODALITAT'
    ]
    resultado = {}
    for i, row in df.iterrows():
        celda_c = str(row[2]).strip() if pd.notna(row[2]) else ''
        for etiqueta in etiquetas:
            if etiqueta.lower() in celda_c.lower():
                valor = str(row[3]).strip() if pd.notna(row[3]) else ''
                # Normalizar fechas si es el campo de fechas
                if etiqueta == 'FECHAS REALIZACIÓN / DATES REALITZACIÓ':
                    valor = normaliza_fechas_realizacion(valor)
                    if valor == -1:
                        return -1
                resultado[etiqueta] = valor
    #print(json.dumps(resultado, ensure_ascii=False, indent=2))

    # Validar formato del código de edición
    codigo = resultado.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')
    if codigo:
        # Formato esperado: 25/26/27 + FP + 2 números + 2 letras + 3 números
        if not re.match(r'^(25|26|27)FP\d{2}[A-Z]{2}\d{3}$', codigo):
            print(
                "Error en format de codi",
                f"El codi d'edició '{codigo}' no compleix el format esperat.\n\n"
                f"Format requerit: 26FP##LL###\n"
                f"Exemple: 26FP43CF123"
            )
            return -1

    return resultado

# Ejemplo de uso:
# extraer_datos_identificativos('FITXA-ECONOMICA.xlsx')


def process_excel(nombre_archivo):
    #status_label.config(text="Processant arxiu Excel...")
    xl = pd.ExcelFile(nombre_archivo)
    hoja = xl.sheet_names[0]
    df = xl.parse(hoja, header=None)
    cabecera = [
        'NOMBRE Y APELLIDOS o EMPRESA / NOM I COGNOMS o EMPRESA',
        'DNI / CIF',
        'JURÍDICO',
        'MINUTA / DIETA / FACTURA/ MATERIAL',
        "TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*",
        'UNIDADES/UNITATS',
        'Solo en caso de tutorización indicar Nº ALUMNANOS/AS TUTORIZADOS',
        'TARIFICACIÓN APLICADA (€)',
        'IMPORTE / IMPORT (€)'
    ]
    # Fila 22 (índice 21), y reseteamos índices
    data_rows = df.iloc[21:, :].reset_index(drop=True)
    header_row = data_rows.iloc[0].fillna('')
    # Detección automática de las columnas de interés
    mapping = {}
    for col in range(len(header_row)):
        cell = str(header_row[col]).strip()
        for campo in cabecera:
            if campo in cell:
                mapping[col] = campo
                break
    datos = []
    for i in range(1, len(data_rows)):
        fila = data_rows.iloc[i]
        entry = {}
        for col, campo in mapping.items():
            entry[campo] = fila[col]
        nombre = str(entry.get(cabecera[0], '')).strip()
        dni = str(entry.get(cabecera[1], '')).strip()
        # Solo añadimos filas con nombre y dni válidos
        if nombre and dni and nombre.lower() != 'nan' and dni.lower() != 'nan':
            datos.append(entry)
    # Agrupamos por persona/empresa
    agrupado = {}
    for entry in datos:
        nombre = str(entry.get(cabecera[0], '')).strip()
        dni = str(entry.get(cabecera[1], '')).strip()
        clave = (nombre, dni)
        if clave not in agrupado:
            agrupado[clave] = []
        # Excluimos campos de nombre y dni internos por síntesis
        entry_light = {k: v for k, v in entry.items() if k not in [cabecera[0], cabecera[1]]}
        agrupado[clave].append(entry_light)
    # Montamos el JSON deseado como lista de personas
    resultado = []
    for (nombre, dni), movimientos in agrupado.items():
        resultado.append({
            "Nombre": nombre,
            "DNI": dni,
            "Movimientos": movimientos
        })
    # Mostrar por pantalla
    # print(json.dumps(resultado, ensure_ascii=False, indent=2))
    return resultado



def generar_certifica_sdgfp(datos, identificativos, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)
    
    imagen_path = "./a.png"
    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(13)

    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    encabezado.add_run("\nJORDI MARTÍ GUIU, SUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT\n \n")

    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False

    fila = tabla.rows[0]
    fila.cells[0].width = Cm(3)
    fila.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    run = fila.cells[0].paragraphs[0].add_run("CERTIFICA")
    run.bold = True

    movimientos = datos['Movimientos']
    dni = str(datos['DNI'])
    nombre = datos['Nombre'].upper()

    partes = []
    for mov in movimientos:
        if mov.get('TARIFICACIÓN APLICADA (€)', 0) == 300:
            tipo_intervencio = mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "").strip().lower()
            if tipo_intervencio == "tutorización":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "tutorització"
                pagament_text = " setmanes a 300 euros/setmana"
            elif tipo_intervencio == "elaboración de casos-actividades prácticas":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "el·laboració de casos-activitats pràctiques"
                pagament_text = " unitats a 300 euros/unitat"
            elif tipo_intervencio == "ponente":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "ponent"
                pagament_text = " hores a 90 euros/hora"
            else:
                pagament_text = " REVISAR TIPO INTERVENCIÓ"
        else:
            pagament_text = f" hores a {str(mov.get('TARIFICACIÓN APLICADA (€)', '')).lower()} euros/hora"
        partes.append(str(mov.get('UNIDADES/UNITATS', '')).lower() + pagament_text + " en concepte de " + str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).lower())

    cadena_movs = ", ".join(partes)
    # Determinar el código de aplicación según el campo 'JURÍDICO'
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        aplicacion = "233.02"
        juridico = "funcionari"
    else:
        aplicacion = "226.06"

    # Calcular la suma total de los importes
    total_importe = sum(
        float(mov.get('IMPORTE / IMPORT (€)', 0) or 0)
        for mov in movimientos
        if mov.get('IMPORTE / IMPORT (€)', '') not in [None, '', 'nan']
    )

    # Añadir el cuerpo con palabras clave en negrita
    parrafo_cuerpo = fila.cells[1].paragraphs[0]


    def add_bold(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True

    def add_normal(paragraph, text):
        paragraph.add_run(text)

    # Construir el cuerpo con formato
    add_normal(parrafo_cuerpo, f"{nombre} amb NIF: {dni.upper()}\n")
    add_normal(parrafo_cuerpo, juridico + ", Designat/da per esta Subdirecció, ha impartit satisfactòriament " + cadena_movs + " del curs amb les següents dades:\n\n")

    if identificativos:
        add_bold(parrafo_cuerpo, "Codi: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Títol: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Data de realització: ")
        add_normal(parrafo_cuerpo, str("del " + identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')).lower() + "\n")
        add_bold(parrafo_cuerpo, "Lloc de realització: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('MODALIDAD/MODALITAT', '')).upper() + "\n\n")

    add_normal(parrafo_cuerpo, "Per la qual cosa, cal fer-li el pagament corresponent per un total de ")
    add_normal(parrafo_cuerpo, str(total_importe) + " euros (")
    add_normal(parrafo_cuerpo, num2words(total_importe, lang='ca').lower())
    add_normal(parrafo_cuerpo, f"), per l’aplicació {aplicacion}")

    fila.cells[1].width = Cm(13)
    # fila.cells[1].text = 
    doc_name = f"{identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')}_CERTIFICA_{datos['Nombre'].replace(' ', '_')}.docx"


    save_path = doc_name
    
    # doc.save(save_path)
    BufferIO = BytesIO()
    doc.save(BufferIO)
    BufferIO.seek(0)
    return BufferIO, save_path


# Generar documento Word
# datos es un diccionario con 'Nombre', 'DNI' y 'Movimientos' (lista de dicts)

def designasdgfp(datos, identificativos, numero_a_letras=lambda x:str(x)):

    print("Generando designa SDGFP para:", datos['Nombre'])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)



    # imagen_path = resource_path('a.png')
    
    imagen_path = "./a.png"
    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(13)

    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    encabezado.add_run("\nJORDI MARTÍ GUIU, SUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT\n \n")

    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False

    fila = tabla.rows[0]
    fila.cells[0].width = Cm(3)
    fila.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    run = fila.cells[0].paragraphs[0].add_run("DESIGNA")
    run.bold = True

    movimientos = datos['Movimientos']
    dni = str(datos['DNI'])
    nombre = datos['Nombre'].upper()

    partes = []
    for mov in movimientos:
        if mov.get('TARIFICACIÓN APLICADA (€)', 0) == 300:
            tipo_intervencio = mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "").strip().lower()
            if tipo_intervencio == "tutorización":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "tutorització"
                pagament_text = " setmanes a 300 euros/setmana"
            elif tipo_intervencio == "elaboración de casos-actividades prácticas":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "el·laboració de casos-activitats pràctiques"
                pagament_text = " unitats a 300 euros/unitat"
            elif tipo_intervencio == "ponente":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "ponent"
                pagament_text = " hores a 90 euros/hora"
            else:
                pagament_text = " REVISAR TIPO INTERVENCIÓ"
        else:
            pagament_text = f" hores a {str(mov.get('TARIFICACIÓN APLICADA (€)', '')).lower()} euros/hora"
        partes.append(str(mov.get('UNIDADES/UNITATS', '')).lower() + pagament_text + " en concepte de " + str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).lower())

    cadena_movs = ", ".join(partes)
    # Determinar el código de aplicación según el campo 'JURÍDICO'
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        aplicacion = "233.02"
        juridico = "funcionari"
    else:
        aplicacion = "226.06"

    # Calcular la suma total de los importes
    total_importe = sum(
        float(mov.get('IMPORTE / IMPORT (€)', 0) or 0)
        for mov in movimientos
        if mov.get('IMPORTE / IMPORT (€)', '') not in [None, '', 'nan']
    )

    cuerpo = (
        f"{nombre} amb nif: {dni.upper()}\n" +
        juridico + ", perquè impartisca " + cadena_movs + " del curs amb les següents dades:\n"
        + ("Codi: " + str(identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')).upper() + "\n" if identificativos else "")
        + ("Títol: " + str(identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')).upper() + "\n\n" if identificativos else "")
        + ("Data de realització: " + str("del " + identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')).lower() + "\n" if identificativos else "")
        + ("Lloc de realització: " + str(identificativos.get('MODALIDAD/MODALITAT', '')).upper() + "\n\n" if identificativos else "")
        + "caldrà pagar-li en concepte d´assistències la quantitat de "
        + str(total_importe) + " euros ("
        + numero_a_letras(total_importe).lower() +
        f"), per l’aplicació {aplicacion}"
    )


    # Añadir el cuerpo con palabras clave en negrita
    parrafo_cuerpo = fila.cells[1].paragraphs[0]


    def add_bold(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True

    def add_normal(paragraph, text):
        paragraph.add_run(text)

    # Construir el cuerpo con formato
    add_normal(parrafo_cuerpo, f"{nombre} amb NIF: {dni.upper()}\n")
    add_normal(parrafo_cuerpo, juridico + ", perquè impartisca " + cadena_movs + " del curs amb les següents dades:\n\n")

    if identificativos:
        add_bold(parrafo_cuerpo, "Codi: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Títol: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Data de realització: ")
        add_normal(parrafo_cuerpo, str("del " + identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')).lower() + "\n")
        add_bold(parrafo_cuerpo, "Lloc de realització: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('MODALIDAD/MODALITAT', '')).upper() + "\n\n")

    add_normal(parrafo_cuerpo, "caldrà pagar-li en concepte d´assistències la quantitat de ")
    add_normal(parrafo_cuerpo, str(total_importe) + " euros (")
    add_normal(parrafo_cuerpo, num2words(total_importe, lang='ca').lower())
    add_normal(parrafo_cuerpo, f"), per l’aplicació {aplicacion}")

    fila.cells[1].width = Cm(13)
    # fila.cells[1].text = 
    doc_name = f"{identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')}_DESIGNA_{datos['Nombre'].replace(' ', '_')}.docx"

    save_path = doc_name
   
    # doc.save(save_path)

    BufferIO = BytesIO()
    doc.save(BufferIO)
    BufferIO.seek(0)
    return BufferIO, save_path


def generar_skills(datos, identificativos, partida, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    imagen_path = './b.png'

    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(10)

    # Encabezado de autoridad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nMarta Armendia Santos, directora general de Formació Professional, de la Conselleria d’Educació, Cultura i Universitats\n").bold = True

    # RESOLUCIÓN
    p_resolc = doc.add_paragraph()
    p_resolc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_resolc = p_resolc.add_run("RESOLC")
    run_resolc.bold = True

    # 1. Designación del personal docente
    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')

    fechas = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')
    modalidad = identificativos.get('MODALIDAD/MODALITAT', '')




    # Determinar si es funcionario GVA
    movimientos = datos['Movimientos']
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        # Si modalidad contiene "online" o "on line" (ignorando mayúsculas/minúsculas), poner "de forma online", si no, poner "presencial a"
        modalidad_lower = modalidad.lower()
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"1. Designar el personal docent que a continuació es relaciona com a formadors, "
            f"per a formar part de l’equip docent que impartirà la formació {codigo} - {curso}, "
            f"{modalidad_text} del {fechas}."
        )
    else:
        modalidad_lower = modalidad.lower()
        nombre = datos.get('Nombre', '')
        dni = datos.get('DNI', '').replace(' ', '')
        # Calcular total de horas (sumar UNIDADES/UNITATS si es relevante)
        total_hores = sum(
            float(mov.get('UNIDADES/UNITATS', 0) or 0)
            for mov in movimientos
            if str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower() == "síncrona"
        )
        # Si no hay horas, dejarlo vacío o poner el total de unidades
        if not total_hores:
            total_hores = sum(float(mov.get('UNIDADES/UNITATS', 0) or 0) for mov in movimientos)
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"1. Designar a les persones que a continuació es relaciona com a formadors, "
            f"per a formar part de l’equip que impartirà la formació {codigo} - {curso}, "
            f"{modalidad_text} del {fechas}."
        )

    doc.add_paragraph(designa_text)

    # TABLA CENTRAL
    movimientos = datos['Movimientos']


    tabla = doc.add_table(rows=1, cols=6)
    # Centrar contenido horizontal y verticalmente en toda la tabla
    # Centrar contenido horizontal y verticalmente en toda la tabla (todas las filas y columnas)
    
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    headers = ["NOM I COGNOMS", "DNI", "UNITATS", "CONCEPTE", "IMPORT PER UNITAT", "TOTAL"]
    for i in range(6):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fondo gris claro para la cabecera
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)

    for mov in movimientos:
        row = tabla.add_row().cells
        row[0].text = datos['Nombre']
        row[1].text = str(datos['DNI'])
        # Añadir sufijo a 'unitats' según el valor de 'concepte'
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        unitats = str(mov.get('UNIDADES/UNITATS', ''))
        if concepte == "síncrona":
            unitats = f"{unitats} hores"
        elif concepte == "elaboración de casos-actividades prácticas":
            unitats = f"{unitats} casos"
        elif concepte == "tutorización":
            unitats = f"{unitats} setmanes"
        row[2].text = unitats
        row[3].text = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", ""))
        # Añadir sufijo a 'tarificació' según el valor de 'concepte'
        tarificacio = str(mov.get('TARIFICACIÓN APLICADA (€)', ''))
        # Traducción de concepte al valenciano
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
            tarificacio = f"{tarificacio} €/hora"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
            tarificacio = f"{tarificacio} €/cas"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
            tarificacio = f"{tarificacio} €/setmana"
        elif concepte == "ponente":
            concepte_val = "ponent"
            tarificacio = f"{tarificacio} €/hora"
        else:
            concepte_val = concepte
        row[3].text = concepte_val
        row[4].text = tarificacio
        row[5].text = str(mov.get('IMPORTE / IMPORT (€)', '')) + " €"

    # Total general
    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)
    row = tabla.add_row().cells
    # Deja las columnas 0 a 4 vacías y pon "TOTAL" en la columna 5
    for i in range(4):
        row[i].text = ""
        # Quitar borde izquierdo y borde inferior
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        # Eliminar borde izquierdo
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:left w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde inferior
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde derecho
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w')))
    
    
    p_total = row[4].paragraphs[0]
    run_total = p_total.add_run("TOTAL")
    run_total.bold = True
    run_total.font.size = Pt(11)
    # Fondo verde claro para la celda "TOTAL"
    tc_total = row[4]._tc
    tcPr_total = tc_total.get_or_add_tcPr()
    shd_total = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_total.append(shd_total)

    p_importe = row[5].paragraphs[0]
    run_importe = p_importe.add_run(str(importe_total)+"€")
    run_importe.bold = True
    run_importe.font.size = Pt(11)
    # Fondo verde claro para la celda de importe total
    tc_importe = row[5]._tc
    tcPr_importe = tc_importe.get_or_add_tcPr()
    shd_importe = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_importe.append(shd_importe)
    
    # Ajustar la altura de la última fila a 2 cm

    tabla.rows[-1].height = Cm(0.8)



    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

    # Traducir y unir los conceptos en valenciano
    conceptos_valenciano = []
    for mov in movimientos:
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
        elif concepte == "ponente":
            concepte_val = "ponent"
        else:
            concepte_val = concepte
        conceptos_valenciano.append(concepte_val)
        # G01090205GE00000.422C00.TE22000053
    doc.add_paragraph(
        "\n2. Aprovar el gasto per un import total de " +
        f"{importe_total} € en concepte de " +
        " i ".join(set(conceptos_valenciano)) +
        ", per la seua participació en l’activitat esmentada i per actuar fora de l’horari normal de treball. "
        f"Este import s’abonarà d’acord amb el Decret 24/1997, d’11 de febrer, i les seues modificacions posteriors, sobre indemnitzacions per raó del servei i gratificacions per serveis extraordinaris, amb càrrec a l’aplicació pressupostària {partida}, del pressupost de la Generalitat Valenciana per a l’any 2026."
    )

    doc.add_paragraph(
        "3. Esta actuació està cofinançada pel Fons Social Europeu i pel Ministeri d’Educació, "
        "Formació Professional i Esports en el marc del programa d’ocupació, formació i educació del període 2021-2027.\n"
    )

    # Pie y firma
    # Pie y firma centrados
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("València, en data i signatura electrònica\n\n\n\n\n")

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("Marta Armendia Santos\nDirectora General de Formació Professional, de la Conselleria d’Educació, Cultura i Universitats")

    # Nombre de archivo y guardado
    doc_name = f"{codigo}_DESIGNA_{datos['Nombre'].replace(' ', '_')}.docx"
  
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return (buffer, doc_name)




def generar_skills_resolc(datos, identificativos, partida, fecha, centre_educatiu, carrec, numero_a_letras=lambda x:str(x)):

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    imagen_path = "./b.png"

    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Times New Roman'
    fuente.size = Pt(12)

    # Encabezado de autoridad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nMARTA ARMENDIA SANTOS, DIRECTORA GENERAL DE FORMACIÓ PROFESSIONAL\n")

    # 1. Designación del personal docente
    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')

    fechas = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')
    modalidad = identificativos.get('MODALIDAD/MODALITAT', '')
    movimientos = datos['Movimientos']
    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)


    # Determinar si es funcionario GVA
    
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        # Si modalidad contiene "online" o "on line" (ignorando mayúsculas/minúsculas), poner "de forma online", si no, poner "presencial a"
        modalidad_lower = modalidad.lower()
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"

        designa_text = (
            f"Vist l'informe del cap de servei del {fecha}, corresponent a la formació {codigo} - {curso} "
            f"realitzada {modalidad_text} del {fechas}.\n"
            f"Vist que els col·laboradors externs han realitzat en els termes establits i de manera adequada la labor "
            f"per a la qual van ser designats."
        )

        texto = (
            f"Que ordene el pagament als col·laboradors externs relacionats a continuació, "   
            f"l'import total de {importe_total} €, amb la distribució indicada, per actuar com a "
            f"col·laboradors en l'activitat de formació {codigo} - {curso}, per actuar fora de l'horari normal de treball i amb càrrec a "
            f"l'aplicació pressupostària {partida}, de conformitat amb el DECRET 80/2025, de 3 de juny, del Consell "
            f"sobre indemnitzacions per raó del servei i gratificacions per serveis extraordinaris."
        )

    else:
        modalidad_lower = modalidad.lower()
        nombre = datos.get('Nombre', '')
        dni = datos.get('DNI', '').replace(' ', '')
        # Calcular total de horas (sumar UNIDADES/UNITATS si es relevante)
        total_hores = sum(
            float(mov.get('UNIDADES/UNITATS', 0) or 0)
            for mov in movimientos
            if str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower() == "síncrona"
        )
        # Si no hay horas, dejarlo vacío o poner el total de unidades
        if not total_hores:
            total_hores = sum(float(mov.get('UNIDADES/UNITATS', 0) or 0) for mov in movimientos)
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"Vist l'informe del cap de servei del {fecha}, corresponent a la formació {codigo} - {curso} "
            f"realitzada {modalidad_text} del {fechas}.\n"
            f"Vist que els professors de la formació han realitzat en els termes establits i de manera adequada la labor "
            f"per a la qual van ser designats."
        )

        texto = (
            f"Que ordene el pagament a les persones relacionades a continuació, "   
            f"l'import total de {importe_total} €, amb la distribució indicada, per actuar com a "
            f"col·laboradors en l'activitat de formació {codigo} - {curso} amb càrrec a "
            f"l'aplicació pressupostària {partida}, de conformitat amb el DECRET 80/2025, de 3 de juny, del Consell "
            f"sobre indemnitzacions per raó del servei i gratificacions per serveis extraordinaris."
        )

    p1 = doc.add_paragraph(designa_text)

    # RESOLUCIÓN
    p_resolc = doc.add_paragraph()
    p_resolc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_resolc = p_resolc.add_run("RESOLC")
    run_resolc.bold = True

    # TABLA CENTRAL
    movimientos = datos['Movimientos']

    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)

    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla = doc.add_table(rows=1, cols=8)
    # Centrar contenido horizontal y verticalmente en toda la tabla
    # Centrar contenido horizontal y verticalmente en toda la tabla (todas las filas y columnas)
    
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    # Ajuste de altura y alineación de la fila de cabecera, y tamaño de fuente 10
    headers = ["NOM I COGNOMS", "DNI","CENTRE EDUCATIU","CÀRREC", "UNITATS", "CONCEPTE", "IMPORT PER UNITAT", "TOTAL"]
    
    for i in range(8):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        run.font.size = Pt(8)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fondo gris claro para la cabecera
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)

    for mov in movimientos:
        row = tabla.add_row().cells
        # Ajustar el tamaño de la fuente por defecto a 10 pt para el contenido de las filas
        row[0].text = datos['Nombre']
        row[1].text = str(datos['DNI'])
        row[2].text = centre_educatiu
        row[3].text = carrec
        # Añadir sufijo a 'unitats' según el valor de 'concepte'
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        unitats = str(mov.get('UNIDADES/UNITATS', ''))
        if concepte == "síncrona":
            unitats = f"{unitats} hores"
        elif concepte == "elaboración de casos-actividades prácticas":
            unitats = f"{unitats} casos"
        elif concepte == "tutorización":
            unitats = f"{unitats} setmanes"
        row[4].text = unitats
        # row[5].text = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", ""))
        # Añadir sufijo a 'tarificació' según el valor de 'concepte'
        tarificacio = str(mov.get('TARIFICACIÓN APLICADA (€)', ''))
        # Traducción de concepte al valenciano
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
            tarificacio = f"{tarificacio} €/hora"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
            tarificacio = f"{tarificacio} €/cas"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
            tarificacio = f"{tarificacio} €/setmana"
        elif concepte == "ponente":
            concepte_val = "ponent"
            tarificacio = f"{tarificacio} €/hora"
        else:
            concepte_val = concepte
        row[5].text = concepte_val
        row[6].text = tarificacio
        row[7].text = str(mov.get('IMPORTE / IMPORT (€)', '')) + " €"
        for cell in row:  # Recorre TODAS las celdas de la fila
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Total general
    
    row = tabla.add_row().cells
    # Deja las columnas 0 a 4 vacías y pon "TOTAL" en la columna 5
    for i in range(6):
        row[i].text = ""
        # Quitar borde izquierdo y borde inferior
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        # Eliminar borde izquierdo
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:left w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde inferior
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde derecho
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w')))
    
    
    p_total = row[6].paragraphs[0]
    run_total = p_total.add_run("TOTAL")
    run_total.bold = True
    run_total.font.size = Pt(11)
    # Fondo verde claro para la celda "TOTAL"
    tc_total = row[6]._tc
    tcPr_total = tc_total.get_or_add_tcPr()
    shd_total = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_total.append(shd_total)

    p_importe = row[7].paragraphs[0]
    run_importe = p_importe.add_run(str(importe_total)+"€")
    run_importe.bold = True
    run_importe.font.size = Pt(11)
    # Fondo verde claro para la celda de importe total
    tc_importe = row[7]._tc
    tcPr_importe = tc_importe.get_or_add_tcPr()
    shd_importe = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_importe.append(shd_importe)
    
    # Ajustar la altura de la última fila a 2 cm

    tabla.rows[-1].height = Cm(0.8)



    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

    # Traducir y unir los conceptos en valenciano
    conceptos_valenciano = []
    for mov in movimientos:
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
        elif concepte == "ponente":
            concepte_val = "ponent"
        else:
            concepte_val = concepte
        conceptos_valenciano.append(concepte_val)
        # G01090205GE00000.422C00.TE22000053


    


    # Pie y firma
    # Pie y firma centrados
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("\nValència, en data i signatura electrònica\n\n\n\n\n")

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("Marta Armendia Santos\nDirectora General de Formació Professional")

    # Nombre de archivo y guardado
    doc_name = f"{codigo}_RESOLC_{datos['Nombre'].replace(' ', '_')}.docx"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return (buffer, doc_name)




def generar_skills_certifica(datos, identificativos, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    imagen_path = "./b.png"

    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(10)

    # Encabezado de autoridad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\nDavid Montalvà Furió, cap del Servei d’Orientació Professional, de la Direcció General de Formació Professional.\n").bold = True

    # RESOLUCIÓN
    p_resolc = doc.add_paragraph()
    p_resolc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_resolc = p_resolc.add_run("INFORMA")
    run_resolc.bold = True

    # 1. Designación del personal docente
    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')

    fechas = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')
    modalidad = identificativos.get('MODALIDAD/MODALITAT', '')




    # Determinar si es funcionario GVA
    movimientos = datos['Movimientos']
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        # Si modalidad contiene "online" o "on line" (ignorando mayúsculas/minúsculas), poner "de forma online", si no, poner "presencial a"
        modalidad_lower = modalidad.lower()
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"Que el personal docent que es relaciona a continuació ha format part com a personal col·laborador per "
            f"a formar part de l'equip docent que van ser anomenats per resolució de la Direcció General de "
            f"Formació Professional, per al curs «{codigo} - {curso}» "
            f"realitzat {modalidad_text} del {fechas}.\n"
        )
    else:
        modalidad_lower = modalidad.lower()
        nombre = datos.get('Nombre', '')
        dni = datos.get('DNI', '').replace(' ', '')
        # Calcular total de horas (sumar UNIDADES/UNITATS si es relevante)
        total_hores = sum(
            float(mov.get('UNIDADES/UNITATS', 0) or 0)
            for mov in movimientos
            if str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower() == "síncrona"
        )
        # Si no hay horas, dejarlo vacío o poner el total de unidades
        if not total_hores:
            total_hores = sum(float(mov.get('UNIDADES/UNITATS', 0) or 0) for mov in movimientos)
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"Que el personal que es relaciona a continuació ha format part com a personal col·laborador per "
            f"a formar part de l'equip docent que van ser anomenats per resolució de la Direcció General de "
            f"Formació Professional, per al curs «{codigo} - {curso}» "
            f"realitzat {modalidad_text} del {fechas}.\n"
        )

    doc.add_paragraph(designa_text)

    # TABLA CENTRAL
    movimientos = datos['Movimientos']


    tabla = doc.add_table(rows=1, cols=6)
    # Centrar contenido horizontal y verticalmente en toda la tabla
    # Centrar contenido horizontal y verticalmente en toda la tabla (todas las filas y columnas)
    
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    headers = ["NOM I COGNOMS", "DNI", "UNITATS", "CONCEPTE", "IMPORT PER UNITAT", "TOTAL"]
    for i in range(6):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fondo gris claro para la cabecera
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)

    for mov in movimientos:
        row = tabla.add_row().cells
        row[0].text = datos['Nombre']
        row[1].text = str(datos['DNI'])
        # Añadir sufijo a 'unitats' según el valor de 'concepte'
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        unitats = str(mov.get('UNIDADES/UNITATS', ''))
        if concepte == "síncrona":
            unitats = f"{unitats} hores"
        elif concepte == "elaboración de casos-actividades prácticas":
            unitats = f"{unitats} unitats"
        elif concepte == "tutorización":
            unitats = f"{unitats} setmanes"
        row[2].text = unitats
        row[3].text = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", ""))
        # Añadir sufijo a 'tarificació' según el valor de 'concepte'
        tarificacio = str(mov.get('TARIFICACIÓN APLICADA (€)', ''))
        # Traducción de concepte al valenciano
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
            tarificacio = f"{tarificacio} €/hora"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
            tarificacio = f"{tarificacio} €/unitat"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
            tarificacio = f"{tarificacio} €/setmana"
        elif concepte == "ponente":
            concepte_val = "ponent"
            tarificacio = f"{tarificacio} €/hora"   
        else:
            concepte_val = concepte
        row[3].text = concepte_val
        row[4].text = tarificacio
        row[5].text = str(mov.get('IMPORTE / IMPORT (€)', '')) + " €"

    # Total general
    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)
    row = tabla.add_row().cells
    # Deja las columnas 0 a 4 vacías y pon "TOTAL" en la columna 5
    for i in range(4):
        row[i].text = ""
        # Quitar borde izquierdo y borde inferior
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        # Eliminar borde izquierdo
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:left w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde inferior
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde derecho
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w')))
    
    
    p_total = row[4].paragraphs[0]
    run_total = p_total.add_run("TOTAL")
    run_total.bold = True
    run_total.font.size = Pt(11)
    # Fondo verde claro para la celda "TOTAL"
    tc_total = row[4]._tc
    tcPr_total = tc_total.get_or_add_tcPr()
    shd_total = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_total.append(shd_total)

    p_importe = row[5].paragraphs[0]
    run_importe = p_importe.add_run(str(importe_total)+"€")
    run_importe.bold = True
    run_importe.font.size = Pt(11)
    # Fondo verde claro para la celda de importe total
    tc_importe = row[5]._tc
    tcPr_importe = tc_importe.get_or_add_tcPr()
    shd_importe = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_importe.append(shd_importe)
    
    # Ajustar la altura de la última fila a 2 cm

    tabla.rows[-1].height = Cm(0.8)



    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

    # Traducir y unir los conceptos en valenciano
    conceptos_valenciano = []
    for mov in movimientos:
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
        else:
            concepte_val = concepte
        conceptos_valenciano.append(concepte_val)

    doc.add_paragraph(
        "\nAquesta actuació està co-finançada per el Fons Social Europeu i per el Ministeri d'Educació, Formació Professional i Esports en el marc del programa d'ocupació, formació i educació 2021-2027.\n "
    )

    # Pie y firma
    # Pie y firma centrados
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("València, en data i signatura electrònica\n\n\n\n\n")

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("David Montalvà Furió\nCap de servei d'Orientació Professional, de la Direcció General de Formació Professional")

    # Nombre de archivo y guardado
    doc_name = f"{codigo}_CERTIFICA_INFORME_POSTERIOR_{datos['Nombre'].replace(' ', '_')}.docx"
    save_path = doc_name

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return (buffer, doc_name)

###### GENERA RESOLC I INFORME SKILLS ######





def show_json(json_data):
    print("Datos JSON:")
    print(json_data)


def crea_minuta_skills_docx(dades, identificativos):
    def crea_docx(datos):
        print("Generando minuta para:", datos.get("Nombre", ""))
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.1)
        section.footer_distance = Cm(0.1)

        
        imagen_path = "./b.png"
        doc.add_picture(imagen_path, width=Cm(15.0))

        estilo = doc.styles['Normal']
        fuente = estilo.font
        fuente.name = 'Calibri'
        fuente.size = Pt(13)

        encabezado = doc.add_paragraph()
        encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
        encabezado.add_run("GRATIFICACIÓ PER ACTIVITAT DOCENT")
        encabezado.runs[0].bold = True
        encabezado.runs[0].font.size = Pt(14)

            # TABLA DE DATOS DEL PERCEPTOR/A
        tabla = doc.add_table(rows=8, cols=6)
        tabla.autofit = True
        for row in tabla.rows:
            row.height = Cm(0.8)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Borde externo doble y sin bordes internos

        tbl = tabla._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.append(tblPr)

        borders = parse_xml(r'''
            <w:tblBorders %s>
            <w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:right w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)

        # Unir las celdas de la primera fila desde la segunda hasta la última
        primera_fila = tabla.rows[0]
        celda_merged = primera_fila.cells[0]
        for i in range(1, len(primera_fila.cells)):
            celda_merged = celda_merged.merge(primera_fila.cells[i])

        run = primera_fila.cells[0].paragraphs[0].add_run("DADES DEL PERCEPTOR/A")
        run.underline = True
        run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        primera_fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       

        segona_fila = tabla.rows[1]
        celda_merged = segona_fila.cells[0]
        for i in range(1, len(segona_fila.cells)):
            celda_merged = celda_merged.merge(segona_fila.cells[i])

        

        segona_fila.height = Cm(0.2)
        segona_fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        

        tercera_fila = tabla.rows[2]
        celda_merged = tercera_fila.cells[0]
        for i in range(1, len(tercera_fila.cells)):
            celda_merged = celda_merged.merge(tercera_fila.cells[i])
        run2 = tercera_fila.cells[0].paragraphs[0].add_run("NOM I COGNOMS: ")
        run2.bold = True
        run3 = tercera_fila.cells[0].paragraphs[0].add_run(str(datos.get("Nombre", "")))
        run3.bold = False


        cuarta_fila = tabla.rows[3]
        celda_merged = cuarta_fila.cells[0]
        for i in range(1, len(cuarta_fila.cells)):
            celda_merged = celda_merged.merge(cuarta_fila.cells[i])

        cuarta_fila.height = Cm(0.2)
        cuarta_fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


        quinta_fila = tabla.rows[4]
        # NIF | <nif> | Email | <email> | Telèfon | <tel>
        #celda_merged = quinta_fila.cells[0].merge(quinta_fila.cells[1])
        celda_merged = quinta_fila.cells[0]
        for i in range(1, len(quinta_fila.cells)):
            celda_merged = celda_merged.merge(quinta_fila.cells[i])


        run_nif = quinta_fila.cells[0].paragraphs[0].add_run("NIF: ")
        run_nif.bold = True
        run_nif2 = quinta_fila.cells[0].paragraphs[0].add_run(str(datos.get("NIF", datos.get("DNI", datos.get("NIF/NIE", ""))))+"  ")
        run_nif2.bold = False
        
        run_email = quinta_fila.cells[0].paragraphs[0].add_run("Email: ")
        run_email.bold = True
        run_email2 = quinta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Email", datos.get("Correo electrónico", datos.get("Correo", "")))).lower() + "  ")
        run_email2.bold = False
        

        run_tel = quinta_fila.cells[0].paragraphs[0].add_run("Telèfon: ")
        run_tel.bold = True
        run_tel2 = quinta_fila.cells[0].paragraphs[0].add_run(str(
            datos.get("Telèfon", datos.get("Teléfono", datos.get("Telefono", "")))
        ))
        run_tel2.bold = False


        sexta_fila = tabla.rows[5]

        celda_merged = sexta_fila.cells[0].merge(sexta_fila.cells[1])
        run_grup = sexta_fila.cells[0].paragraphs[0].add_run("GRUP: ")
        run_grup.bold = True
        run_grup2 = sexta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Grup", "")))
        run_grup2.bold = False


        celda_merged = sexta_fila.cells[2].merge(sexta_fila.cells[3])
        run_nivell = sexta_fila.cells[2].paragraphs[0].add_run("NIVELL: ")
        run_nivell.bold = True
        run_nivell2 = sexta_fila.cells[2].paragraphs[0].add_run(str(datos.get("Nivell", "")))
        run_nivell2.bold = False

        celda_merged = sexta_fila.cells[4].merge(sexta_fila.cells[5])
        run_relacion = sexta_fila.cells[4].paragraphs[0].add_run("RELACIÓ JURÍDICA: ")
        run_relacion.bold = True
        run_relacion2 = sexta_fila.cells[4].paragraphs[0].add_run(str(datos.get("Relacio_juridica", "")))
        run_relacion2.bold = False

        sexta_fila.height = Cm(1.3)
        sexta_fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        septima_fila = tabla.rows[6]

        celda_merged = septima_fila.cells[0]
        for i in range(1, len(septima_fila.cells)):
            celda_merged = celda_merged.merge(septima_fila.cells[i])

        
        run_domicilio = septima_fila.cells[0].paragraphs[0].add_run("DOMICILI: ")
        run_domicilio.bold = True
        run_domicilio2 = septima_fila.cells[0].paragraphs[0].add_run(str(datos.get("Domicili", "")))
        run_domicilio2.bold = False


        # Octava fila: CP, POBLACIÓ, PROVÍNCIA

        octava_fila = tabla.rows[7]

        run_cp = octava_fila.cells[0].paragraphs[0].add_run("CP: ")
        run_cp.bold = True
        run_cp2 = octava_fila.cells[0].paragraphs[0].add_run(str(datos.get("CP", "")))
        run_cp2.bold = False

        celda_merged = octava_fila.cells[1].merge(octava_fila.cells[2])
        run_poblacio = octava_fila.cells[1].paragraphs[0].add_run("POBLACIÓ: ")
        run_poblacio.bold = True
        run_poblacio2 = octava_fila.cells[1].paragraphs[0].add_run(str(datos.get("Población", "")))
        run_poblacio2.bold = False

        celda_merged = octava_fila.cells[4].merge(octava_fila.cells[5])
        run_provincia = octava_fila.cells[4].paragraphs[0].add_run("PROVÍNCIA: ")
        run_provincia.bold = True
        run_provincia2 = octava_fila.cells[4].paragraphs[0].add_run(str(datos.get("Provincia", "")))
        run_provincia2.bold = False

        #####

        doc.add_paragraph("")

        ####

        # TABLA DE DATOS ECONÓMICOS DEL PERCEPTOR/A
        tabla3 = doc.add_table(rows=6, cols=6)
        # tabla.autofit = True
        for row in tabla3.rows:
            row.height = Cm(0.8)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Borde externo doble y sin bordes internos

        tbl = tabla3._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.append(tblPr)

        borders = parse_xml(r'''
            <w:tblBorders %s>
            <w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:right w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)

        # Unir las celdas de la primera fila desde la segunda hasta la última
        primera_fila = tabla3.rows[0]
        celda_merged = primera_fila.cells[0]
        for i in range(1, len(primera_fila.cells)):
            celda_merged = celda_merged.merge(primera_fila.cells[i])

        run = primera_fila.cells[0].paragraphs[0].add_run("DADES ECONÒMIQUES")
        run.underline = True
        run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        primera_fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       

        segona_fila = tabla3.rows[1]
        run2 = segona_fila.cells[0].paragraphs[0].add_run("DESCRIPCIÓ DE L’ACTIVITAT: ")
        run2.bold = True

        celda_merged = segona_fila.cells[0]
        for i in range(0, len(segona_fila.cells)):
            celda_merged = celda_merged.merge(segona_fila.cells[i])

        tercera_fila = tabla3.rows[2]
        tercera_fila.cells[1].paragraphs[0].text = datos["Nombre del curso"]

        celda_merged = tercera_fila.cells[0]
        for i in range(0, len(tercera_fila.cells)):
            celda_merged = celda_merged.merge(tercera_fila.cells[i])

        cuarta_fila = tabla3.rows[3]

        for i in range(0, len(cuarta_fila.cells)):
            celda_merged = cuarta_fila.cells[0].merge(cuarta_fila.cells[i])

        cuarta_fila.cells[0].paragraphs[0].add_run("DATES: ")
        cuarta_fila.cells[0].paragraphs[0].runs[0].bold = True
        run2 = cuarta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Dates_inici_final", "")))
        run2.bold = False


        quinta_fila = tabla3.rows[4]
        for i in range(0, len(quinta_fila.cells)):
            celda_merged = quinta_fila.cells[0].merge(quinta_fila.cells[i])

        
        run = quinta_fila.cells[0].paragraphs[0].add_run("IMPORT BRUT: ")
        run.bold = True
        run2 = quinta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Importe bruto", ""))+ " €")
        run2.bold = False

        sexta_fila = tabla3.rows[5]
        for i in range(0, len(sexta_fila.cells)):
            celda_merged = sexta_fila.cells[0].merge(sexta_fila.cells[i])

        
        sexta_fila.cells[0].paragraphs[0].add_run("IMPORT NET: ")
        sexta_fila.cells[0].paragraphs[0].runs[0].bold = True
        run2 = sexta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Importe neto", "")) + " €")
        run2.bold = False
        



        ###########################################

        doc.add_paragraph("")
        

        ####

        # TABLA DE DATOS bancarios DEL PERCEPTOR/A
        tabla4 = doc.add_table(rows=3, cols=6)
        # tabla.autofit = True
        for row in tabla4.rows:
            row.height = Cm(0.9)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Borde externo doble y sin bordes internos

        tbl = tabla4._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.append(tblPr)

        borders = parse_xml(r'''
            <w:tblBorders %s>
            <w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:right w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)

        # Unir las celdas de la primera fila desde la segunda hasta la última
        primera_fila = tabla4.rows[0]
        celda_merged = primera_fila.cells[0]
        for i in range(1, len(primera_fila.cells)):
            celda_merged = celda_merged.merge(primera_fila.cells[i])

        run = primera_fila.cells[0].paragraphs[0].add_run("DADES BANCÀRIES")
        run.underline = True
        run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        primera_fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       

        segona_filatt4 = tabla4.rows[1]
        celda_merged2 = segona_filatt4.cells[0]
        for i in range(1, len(segona_filatt4.cells)):
            celda_merged2 = celda_merged2.merge(segona_filatt4.cells[i])
   
        run = segona_filatt4.cells[0].paragraphs[0].add_run("IBAN: ")
        run.bold = True
        
        # Formatear IBAN: separar en grupos de 4 caracteres
        iban_raw = str(datos.get("IBAN", "")).replace(" ", "").upper()
        iban_formatted = " ".join([iban_raw[i:i+4] for i in range(0, len(iban_raw), 4)])
        run2 = segona_filatt4.cells[0].paragraphs[0].add_run(iban_formatted)
        run2.bold = False


        # segona_fila.cells[1].paragraphs[0].text = str(datos.get("IBAN", ""))

        tercera_filatt4 = tabla4.rows[2]
        celda_merged = tercera_filatt4.cells[0]
        for i in range(0, len(tercera_filatt4.cells)):
            celda_merged = tercera_filatt4.cells[0].merge(tercera_filatt4.cells[i])

        run = tercera_filatt4.cells[0].paragraphs[0].add_run("BIC: ")
        run.bold = True
        run2 = tercera_filatt4.cells[0].paragraphs[0].add_run(str(datos.get("BIC", "")))
        run2.bold = False

        ####################################

        doc.add_paragraph("Declare que he realitzat la citada activitat en la data que s'assenyala.")

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("(firma digital)")

        # Pie de página
        footer = doc.sections[0].footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_paragraph.add_run("CONSELLERIA D'EDUCACIÓ, CULTURA I UNIVERSITATS\n")
        run.font.size = Pt(7)
        run = footer_paragraph.add_run("Av. Campanar, 32. 46015 - València. CIF S4611001A\n")
        run.font.size = Pt(7)
        run = footer_paragraph.add_run("DIRECCIÓ GENERAL DE FORMACIÓ PROFESSIONAL")
        run.font.size = Pt(7)



        #####################################
        # fila.cells[1].text = 
        codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')
        doc_name = f"{codigo}_MINUTA_{datos['Nombre'].replace(' ', '_')}.docx"


        Buffer = BytesIO()
        doc.save(Buffer)
        Buffer.seek(0)
        return (Buffer, doc_name)
    print ("DADES:", dades)
    Buffer, doc_name = crea_docx(dades)
    print ("DOC NAME:", doc_name)
    return (Buffer, doc_name)


def on_process(json_data, hoja_excel, tipo, resultados=None, minuta_datos=None):
        
        fecha = ""
        centre_educatiu = ""
        carrec = ""        
        
        global t
        t = tipo
        personas = []
        try:
            if t == "min":
                Buffer, path = crea_minuta_skills_docx(dades=minuta_datos, identificativos=hoja_excel)
                return Buffer, path
            elif t == "resolc" or t == "cer" or t == "des" or t == "dessdgfp" or t == "cersdgfp":
                # print("JSON DATA:", json_data)
                for persona in json_data:
                    # Filtrar movimientos que contienen "minuta" en el campo 'MINUTA / DIETA / FACTURA/ MATERIAL'
                    # Solo generar documento si TODOS los movimientos son 'minuta'
                    # Si algún movimiento es "caso-actividad", lo convertimos a "minuta"
                    for mov in persona.get('Movimientos', []):
                        tipo = str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower()
                        if tipo == 'caso-actividad':
                            mov['MINUTA / DIETA / FACTURA/ MATERIAL'] = 'minuta'
                    if all(
                        str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta'
                        for mov in persona.get('Movimientos', [])
                    ):
                        # generar_documento(datos=persona, identificativos=hoja_excel)
                        personas.append(persona)
                        if t == "des":
                            pass
                        elif t == "cer":
                            pass
                        elif t == "resolc":
                            # rea_ventana_fechas(persona.get('Nombre', ''))
                            pass
                
                if t == "des":
                    buffers = []
                    for persona in personas:
                        buffer, path = generar_skills(datos=persona, identificativos=hoja_excel, partida="G01090205GE00000.422C00.22699 fons TE22000053")
                        buffers.append((buffer, path))
                    return (buffers)
                elif t == "cer":
                    buffers = []
                    for persona in personas:
                        buffer, path = generar_skills_certifica(datos=persona, identificativos=hoja_excel)
                        buffers.append((buffer, path))
                    return (buffers)
                elif t == "resolc":
                    # rea_ventana_fechas(persona.get('Nombre', ''))
                    buffers3 = []
                    fecha = ""
                    centre_educatiu = ""
                    carrec = ""
                    for persona in personas:
                        print ("PERSONA:", persona['Nombre'])
                        print ("RESULTADOS:", resultados[0])
                        for res in resultados:
                            print (res['persona'], persona['Nombre'])
                            if res['persona'] == persona['Nombre']:
                                fecha = res['fecha']
                                centre_educatiu = res['centro']
                                carrec = res['cargo']
                        print ("FECHA:", fecha)
                        print ("CENTRE EDUCATIU:", centre_educatiu)
                        print ("CARREC:", carrec)
                        fecha_obj = datetime.strptime(fecha, "%Y-%m-%d")  # ajusta el formato de entrada
                        formato_deseado = fecha_obj.strftime("%d/%m/%Y")
                        buffer, path = generar_skills_resolc(datos=persona, identificativos=hoja_excel, fecha=formato_deseado,  centre_educatiu=centre_educatiu, carrec=carrec,partida="G01090205GE00000.422C00.22699 fons OT23000000")
                        print (path)
                        print (buffer)
                        buffers3.append((buffer, path))
                    return (buffers3)
                elif t == "dessdgfp":
                    buffers4 = []
                    for persona in personas:
                        buffer, path = designasdgfp(datos=persona, identificativos=hoja_excel)
                        buffers4.append((buffer, path))
                    return (buffers4)
                elif t == "cersdgfp":
                    buffers5 = []
                    for persona in personas:
                        buffer, path = generar_certifica_sdgfp(datos=persona, identificativos=hoja_excel)
                        buffers5.append((buffer, path))
                    return (buffers5)
            
        except Exception as e:
            pass
            # status_label.config(text=f"Error: {e}")
            # messagebox.showerror("Error", f"Error processant l'archivo: {e}")

